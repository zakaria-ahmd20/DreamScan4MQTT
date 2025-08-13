from docx import Document

# Create a new Word document
doc = Document()
doc.add_heading('Dream4MQTT – MQTT Broker Vulnerability Scanner', 0)

doc.add_paragraph("Dream4MQTT is a Python-based, modular vulnerability scanner for MQTT brokers. It fingerprints the broker version, checks authentication rules, retrieves CVEs, and generates a detailed vulnerability report in PDF format using LaTeX.")

# Features
doc.add_heading('Features', level=1)
doc.add_paragraph("- Detects MQTT broker version and CPE")
doc.add_paragraph("- Retrieves relevant CVEs from local or online databases")
doc.add_paragraph("- Performs authentication bypass checks")
doc.add_paragraph("- Modular test architecture (easily extensible)")
doc.add_paragraph("- Generates a detailed PDF vulnerability report via LaTeX")

# Requirements
doc.add_heading('Requirements', level=1)
doc.add_heading('Python Dependencies', level=2)
doc.add_paragraph("Install required Python libraries:")
doc.add_paragraph("pip3 install paho-mqtt rule-engine")

doc.add_heading('LaTeX (for PDF report generation)', level=2)
doc.add_paragraph("Dream4MQTT uses PdfLaTeX to compile the vulnerability report.\nOn Linux (Debian/Ubuntu/Kali), run:")
doc.add_paragraph("""sudo apt update
sudo apt install texlive-latex-base texlive-latex-recommended texlive-latex-extra texlive-fonts-recommended texlive-fonts-extra""", style='Code')

doc.add_heading('Testing the Scanner', level=1)
doc.add_paragraph("To quickly test Dream4MQTT, you can run a Mosquitto broker using the pre-built Docker image:")
doc.add_paragraph("""# Pull the Docker image
docker pull zahmed70/mosquitto-1.3.1:latest

# Run the broker on default MQTT port 1883
docker run -it -p 1883:1883 -v ~/mosquitto/config:/mosquitto/config zahmed70/mosquitto-1.3.1""", style='Code')

doc.add_paragraph("Once the broker is running, you can scan it with Dream4MQTT:")
doc.add_paragraph("python3 main.py")

# Example output
doc.add_heading('Example Output', level=1)
doc.add_paragraph("""[*] Scanning MQTT broker and retrieving CVEs...
Raw Broker Version: mosquitto version 2.0.16
Detected Version: 2.0.16
CPE: cpe:2.3:a:eclipse:mosquitto:2.0.16:*:*:*:*:*:*:*
[✓] LaTeX source written to reports/mqtt_cve_report.tex
[*] Compiling PDF...
[✓] Report generated: reports/mqtt_cve_report.pdf""")

# Troubleshooting
doc.add_heading('Troubleshooting', level=1)
doc.add_paragraph("- Missing LaTeX packages: install texlive-latex-extra if you see errors")
doc.add_paragraph("- PDF compilation timeout: increase timeout in generate_report.py")
doc.add_paragraph("- Python module not found (rule_engine or paho-mqtt): pip3 install paho-mqtt rule-engine")

# Notes
doc.add_heading('Notes', level=1)
doc.add_paragraph("- Dream4MQTT is designed for research and testing on brokers you own or have permission to scan.")
doc.add_paragraph("- The scanner is modular; you can add new tests or rules in the modules folder.")

# License
doc.add_heading('License', level=1)
doc.add_paragraph("MIT License")

# Save document
doc.save('/home/seed/Desktop/Dream4MQTT/README.docx')
